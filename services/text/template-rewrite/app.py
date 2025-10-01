from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import List, Dict, Any
import uvicorn
import requests
import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import uuid
from datetime import datetime

app = FastAPI(title="Template Rewrite Service", version="1.0.0")

# Ollama configuration
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")

class ImzaKisi(BaseModel):
    isim: str
    unvan: str

class GerekceRequest(BaseModel):
    konu: str
    icerik_konusu: str
    imza_atacaklar: List[ImzaKisi]

class BelgenetRequest(BaseModel):
    konu: str
    icerik_konusu: str
    imza_atacaklar: List[ImzaKisi]
    format_type: str = "belgenet"  # "gerekce" veya "belgenet"

class TemplateResponse(BaseModel):
    success: bool
    message: str
    file_path: str
    filename: str

@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "template-rewrite"}

def call_ollama_for_content(prompt: str, model: str = "gemma3:27b") -> str:
    """Call Ollama to generate content"""
    try:
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.7,
                    "top_p": 0.9
                }
            },
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            return result.get("response", "")
        else:
            print(f"Ollama API error: {response.status_code}")
            return ""
            
    except Exception as e:
        print(f"Error calling Ollama: {e}")
        return ""

def extract_text_from_word(word_file_path):
    """Extract text from Word document"""
    try:
        doc = Document(word_file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading Word file {word_file_path}: {e}")
        return ""

def parse_gerekce_from_text(text):
    """Parse gerekce content from text"""
    lines = text.split('\n')
    
    # Başlığı bul (genellikle ilk satır)
    title = lines[0] if lines else "Gerekçe Belgesi"
    
    # İçeriği bul (başlık sonrası, imza öncesi)
    content_lines = []
    signature_start = -1
    
    for i, line in enumerate(lines[1:], 1):
        # İmza bölümünü tespit et (isim ve unvan içeren satırlar)
        if any(keyword in line.lower() for keyword in ['müdür', 'başkan', 'mühendis', 'uzman', 'şef']):
            signature_start = i
            break
        content_lines.append(line)
    
    content = '\n'.join(content_lines).strip()
    
    # İmzaları bul
    signatures = []
    if signature_start > 0:
        for line in lines[signature_start:]:
            if line.strip() and not any(keyword in line.lower() for keyword in ['imza', 'tarih', 'onay']):
                # İsim ve unvanı ayır
                parts = line.strip().split()
                if len(parts) >= 2:
                    # Son kelime unvan olabilir
                    name = ' '.join(parts[:-1])
                    title_part = parts[-1]
                    signatures.append({"name": name, "title": title_part})
    
    return {
        "title": title,
        "content": content,
        "signatures": signatures
    }

def parse_belgenet_from_text(text: str) -> Dict[str, Any]:
    """Parse belgenet text into structured format"""
    try:
        # Belgenet formatını parse et: Konu: ... İçerik: ... şeklinde
        documents = []
        
        # Metni bölümlere ayır (--- ile ayrılmış)
        sections = text.split('---')
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
                
            # Konu ve İçerik kısımlarını bul
            if 'Konu:' in section and 'İçerik:' in section:
                konu_part = section.split('Konu:')[1].split('İçerik:')[0].strip()
                icerik_part = section.split('İçerik:')[1].strip()
                
                documents.append({
                    'konu': konu_part,
                    'icerik': icerik_part
                })
        
        if documents:
            return {
                'title': 'Belgenet Evrak Örneği',
                'documents': documents
            }
        
        return None
    except Exception as e:
        print(f"❌ Belgenet parse hatası: {e}")
        return None

def load_template_examples():
    """Load template examples from Word and JSON files"""
    examples = []
    
    # Gerekçe şablonları
    gerekceler_dir = "/app/templates/gerekceler"
    if os.path.exists(gerekceler_dir):
        for filename in os.listdir(gerekceler_dir):
            file_path = os.path.join(gerekceler_dir, filename)
            
            if filename.endswith('.json'):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        example = json.load(f)
                        examples.append(example)
                except Exception as e:
                    print(f"Error loading JSON template {filename}: {e}")
            
            elif filename.endswith('.docx'):
                try:
                    # Word dosyasını oku ve JSON'a çevir
                    text = extract_text_from_word(file_path)
                    if text:
                        example = parse_gerekce_from_text(text)
                        examples.append(example)
                        print(f"Loaded Word template: {filename}")
                except Exception as e:
                    print(f"Error loading Word template {filename}: {e}")
    
    # Belgenet şablonları
    belgenet_dir = "/app/templates/belgenet"
    if os.path.exists(belgenet_dir):
        for filename in os.listdir(belgenet_dir):
            file_path = os.path.join(belgenet_dir, filename)
            
            if filename.endswith('.json'):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        example = json.load(f)
                        examples.append(example)
                except Exception as e:
                    print(f"Error loading JSON template {filename}: {e}")
            
            elif filename.endswith('.docx'):
                try:
                    # Word dosyasını oku ve JSON'a çevir
                    text = extract_text_from_word(file_path)
                    if text:
                        example = parse_belgenet_from_text(text)
                        examples.append(example)
                        print(f"Loaded Word template: {filename}")
                except Exception as e:
                    print(f"Error loading Word template {filename}: {e}")
    
    return examples

def generate_belgenet_content(konu: str, icerik_konusu: str) -> tuple:
    """Generate belgenet content and title using LLM with template examples"""
    
    # Örnek şablonları yükle
    examples = load_template_examples()
    
    # Örnekleri prompt'a ekle
    examples_text = ""
    if examples:
        examples_text = "\n\nÖrnek belgenet şablonları:\n"
        for i, example in enumerate(examples[:2], 1):  # İlk 2 örneği kullan
            examples_text += f"\nÖrnek {i}:\n"
            if 'documents' in example:
                for j, doc in enumerate(example['documents'][:3], 1):  # İlk 3 belge
                    examples_text += f"  Belge {j}:\n"
                    examples_text += f"    Konu: {doc.get('konu', '')}\n"
                    examples_text += f"    İçerik: {doc.get('icerik', '')[:200]}...\n"
    
    prompt = f"""
    Belgenet evrak belgesi oluştur (resmi yazı formatında):
    
    Mevcut Konu: {konu}
    İçerik Konusu: {icerik_konusu}
    {examples_text}
    
    Önce içeriğe uygun yeni bir başlık oluştur, sonra belgenet formatında TEK evrak yaz.
    
    Başlık formatı: "YENİ BAŞLIK: [başlık buraya]"
    Dosya adı formatı: "DOSYA ADI: [2-3 kelimelik kısa isim]"
    
    Belgenet evrak formatı (resmi yazı):
    - Sadece İÇERİK kısmını yaz (Konu: kısmını yazma, zaten var)
    - Resmi yazı dili kullan ("Bilgilerinizi ve gereğini arz ederim" ile bitir)
    - Türkçe yaz
    - Kısa ve öz olsun (2-3 paragraf)
    - Direkt içeriği yaz, "İçerik:" etiketi kullanma
    - Sayı, Tarih, Gönderen, Dağıtım gibi metadata'ları YAZMA
    - Sadece ana içerik paragraflarını yaz
    
    Örnek format:
    [Resmi yazı içeriği - 2-3 paragraf, "Bilgilerinizi ve gereğini arz ederim" ile bitir]
    """
    
    content = call_ollama_for_content(prompt)
    
    # Başlık ve dosya adını ayır
    new_title = konu
    new_filename = "belgenet"
    
    if "YENİ BAŞLIK:" in content:
        parts = content.split("YENİ BAŞLIK:")
        if len(parts) >= 2:
            # Başlık kısmını al ve köşeli parantezleri temizle
            title_part = parts[1].split('\n')[0].strip()
            title_part = title_part.replace('[', '').replace(']', '').strip()
            new_title = title_part
            content = parts[1].split('\n', 1)[1].strip() if len(parts[1].split('\n')) > 1 else parts[1].strip()
    
    if "DOSYA ADI:" in content:
        parts = content.split("DOSYA ADI:")
        if len(parts) >= 2:
            # Dosya adı kısmını al ve köşeli parantezleri temizle
            filename_part = parts[1].split('\n')[0].strip()
            filename_part = filename_part.replace('[', '').replace(']', '').strip()
            new_filename = filename_part
            content = parts[1].split('\n', 1)[1].strip() if len(parts[1].split('\n')) > 1 else parts[1].strip()
    
    return new_title, content, new_filename

def generate_gerekce_content(konu: str, icerik_konusu: str) -> tuple:
    """Generate gerekce content and title using LLM with template examples"""
    
    # Örnek şablonları yükle
    examples = load_template_examples()
    
    # Örnekleri prompt'a ekle
    examples_text = ""
    if examples:
        examples_text = "\n\nÖrnek gerekçe şablonları:\n"
        for i, example in enumerate(examples[:2], 1):  # İlk 2 örneği kullan
            examples_text += f"\nÖrnek {i}:\n"
            examples_text += f"Başlık: {example.get('title', '')}\n"
            examples_text += f"İçerik: {example.get('content', '')[:200]}...\n"
    
    prompt = f"""
    Gerekçe belgesi oluştur:
    
    Mevcut Konu: {konu}
    İçerik Konusu: {icerik_konusu}
    {examples_text}
    
    Önce içeriğe uygun yeni bir başlık oluştur, sonra gerekçe metnini yaz.
    
    Başlık formatı: "YENİ BAŞLIK: [başlık buraya]"
    Dosya adı formatı: "DOSYA ADI: [2-3 kelimelik kısa isim]"
    
    Gerekçe metni için:
    - Başlık kullanma, sadece paragraflar yaz
    - Giriş paragrafı (konunun önemini açıkla)
    - Mevcut durum analizi
    - İhtiyaç ve gerekçe
    - Beklenen faydalar
    - Sonuç ve öneri
    
    Her paragraf 2-3 cümle olsun.
    Resmi dil kullan.
    Türkçe yaz.
    Örnek şablonlardaki yapıyı ve tarzı takip et.
    """
    
    content = call_ollama_for_content(prompt)
    
    # Başlık ve dosya adını ayır
    new_title = konu
    new_filename = "gerekce"
    
    if "YENİ BAŞLIK:" in content:
        parts = content.split("YENİ BAŞLIK:")
        if len(parts) >= 2:
            # Başlık kısmını al ve köşeli parantezleri temizle
            title_part = parts[1].split('\n')[0].strip()
            title_part = title_part.replace('[', '').replace(']', '').strip()
            new_title = title_part
            content = parts[1].split('\n', 1)[1].strip() if len(parts[1].split('\n')) > 1 else parts[1].strip()
    
    if "DOSYA ADI:" in content:
        parts = content.split("DOSYA ADI:")
        if len(parts) >= 2:
            # Dosya adı kısmını al ve köşeli parantezleri temizle
            filename_part = parts[1].split('\n')[0].strip()
            filename_part = filename_part.replace('[', '').replace(']', '').strip()
            new_filename = filename_part
            content = parts[1].split('\n', 1)[1].strip() if len(parts[1].split('\n')) > 1 else parts[1].strip()
    
    return new_title, content, new_filename

def format_signatures(imza_atacaklar: List[ImzaKisi]) -> List[str]:
    """Format signatures for Word document with proper centering"""
    if len(imza_atacaklar) == 1:
        return [f"{imza_atacaklar[0].isim}", f"{imza_atacaklar[0].unvan}"]
    elif len(imza_atacaklar) == 2:
        # İki kişi için eşit aralıklarla ortala
        return [
            f"{imza_atacaklar[0].isim:^30} {imza_atacaklar[1].isim:^30}",
            f"{imza_atacaklar[0].unvan:^30} {imza_atacaklar[1].unvan:^30}"
        ]
    elif len(imza_atacaklar) == 3:
        # Üç kişi için eşit aralıklarla ortala - ünvanlar isimlere göre ortalanmış
        return [
            f"{imza_atacaklar[0].isim:^25} {imza_atacaklar[1].isim:^25} {imza_atacaklar[2].isim:^25}",
            f"{imza_atacaklar[0].unvan:^25} {imza_atacaklar[1].unvan:^25} {imza_atacaklar[2].unvan:^25}"
        ]
    else:
        # 4+ kişi için 2'li satırlar halinde
        lines = []
        for i in range(0, len(imza_atacaklar), 2):
            if i + 1 < len(imza_atacaklar):
                lines.append(f"{imza_atacaklar[i].isim:^30} {imza_atacaklar[i+1].isim:^30}")
                lines.append(f"{imza_atacaklar[i].unvan:^30} {imza_atacaklar[i+1].unvan:^30}")
            else:
                lines.append(f"{imza_atacaklar[i].isim:^60}")
                lines.append(f"{imza_atacaklar[i].unvan:^60}")
        return lines

def create_belgenet_word_document(konu: str, content: str, imza_atacaklar: List[ImzaKisi], filename_base: str = "belgenet") -> str:
    """Create Word document for belgenet - simplified official memo format"""
    doc = Document()
    
    # Konu başlığı (sol üst)
    konu_para = doc.add_paragraph(f"Konu : {konu}")
    konu_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # DAĞITIM YERLERİNE (merkezi)
    dagitim_para = doc.add_paragraph("DAĞITIM YERLERİNE")
    dagitim_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dagitim_para.runs[0].bold = True
    
    # Boşluk
    doc.add_paragraph()
    
    # İçerik - Belgenet formatında (TEK evrak)
    # Direkt içerik paragrafları (resmi yazı formatı)
    for paragraph_text in content.split('\n'):
        if paragraph_text.strip():
            para = doc.add_paragraph(paragraph_text.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Inches(0.5)
    
    # İmza alanı YOK - belgenet formatında imza yok
    
    # Dosya adı oluştur - LLM'den gelen kısa isim
    safe_filename = "".join(c for c in filename_base if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_filename = safe_filename.replace(' ', '_')[:30]  # Maksimum 30 karakter
    filename = f"{safe_filename}_{uuid.uuid4().hex[:8]}.docx"
    filepath = f"/app/outputs/{filename}"
    
    # Kaydet
    doc.save(filepath)
    
    return filepath, filename

def create_gerekce_word_document(konu: str, content: str, imza_atacaklar: List[ImzaKisi], filename_base: str = "gerekce") -> str:
    """Create Word document for gerekce"""
    doc = Document()
    
    # Başlık
    title = doc.add_heading(konu, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Boşluk
    doc.add_paragraph()
    
    # İçerik
    content_paragraph = doc.add_paragraph(content)
    content_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # İmza için boşluklar
    for _ in range(3):
        doc.add_paragraph()
    
    # İmzalar
    signature_lines = format_signatures(imza_atacaklar)
    for line in signature_lines:
        sig_para = doc.add_paragraph(line)
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dosya adı oluştur - LLM'den gelen kısa isim
    safe_filename = "".join(c for c in filename_base if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_filename = safe_filename.replace(' ', '_')[:30]  # Maksimum 30 karakter
    filename = f"{safe_filename}_{uuid.uuid4().hex[:8]}.docx"
    filepath = f"/app/outputs/{filename}"
    
    # Kaydet
    doc.save(filepath)
    
    return filepath, filename

@app.post("/generate-document", response_model=TemplateResponse)
async def generate_document(request: BelgenetRequest):
    """Generate document (gerekce or belgenet format)"""
    try:
        if request.format_type == "belgenet":
            # Belgenet formatında içerik oluştur
            title, content, filename_base = generate_belgenet_content(
                request.konu, 
                request.icerik_konusu
            )
            
            # Belgenet Word belgesi oluştur
            filepath, filename = create_belgenet_word_document(
                title, 
                content, 
                request.imza_atacaklar,
                filename_base
            )
            
            return TemplateResponse(
                success=True,
                message="Belgenet evrak belgesi başarıyla oluşturuldu",
                file_path=filepath,
                filename=filename
            )
        else:
            # Gerekçe formatında içerik oluştur
            title, content, filename_base = generate_gerekce_content(
                request.konu, 
                request.icerik_konusu
            )
            
            # Gerekçe Word belgesi oluştur
            filepath, filename = create_gerekce_word_document(
                title, 
                content, 
                request.imza_atacaklar,
                filename_base
            )
            
            return TemplateResponse(
                success=True,
                message="Gerekçe belgesi başarıyla oluşturuldu",
                file_path=filepath,
                filename=filename
            )
        
    except Exception as e:
        print(f"❌ Belge oluşturma hatası: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-gerekce", response_model=TemplateResponse)
async def generate_gerekce(request: GerekceRequest):
    """Generate gerekce document"""
    try:
        # LLM ile içerik, başlık ve dosya adı oluştur
        title, content, filename_base = generate_gerekce_content(request.konu, request.icerik_konusu)
        
        if not content:
            raise HTTPException(status_code=500, detail="İçerik oluşturulamadı")
        
        # Word belgesi oluştur
        filepath, filename = create_gerekce_word_document(
            title, 
            content, 
            request.imza_atacaklar,
            filename_base
        )
        
        return TemplateResponse(
            success=True,
            message="Gerekçe belgesi başarıyla oluşturuldu",
            file_path=filepath,
            filename=filename
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/templates")
async def get_templates():
    """Get available templates"""
    return {
        "templates": [
            {
                "id": "gerekceler",
                "name": "Gerekçeler",
                "description": "Resmi gerekçe belgesi şablonu",
                "fields": [
                    {"name": "konu", "type": "text", "required": True, "label": "Konu"},
                    {"name": "icerik_konusu", "type": "textarea", "required": True, "label": "İçerik Konusu"},
                    {"name": "imza_atacaklar", "type": "array", "required": True, "label": "İmza Atacaklar"}
                ]
            }
        ]
    }

@app.get("/download/{filename}")
async def download_file(filename: str):
    """Download generated file"""
    file_path = f"/app/outputs/{filename}"
    if os.path.exists(file_path):
        return {"file_path": file_path, "filename": filename}
    else:
        raise HTTPException(status_code=404, detail="Dosya bulunamadı")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8005)